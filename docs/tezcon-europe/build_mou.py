from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("TezCon-Europe-Memorandum-of-Understanding-Draft.docx")

INK = RGBColor(31, 35, 40)
BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 101, 112)
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run(run, size=11, bold=False, italic=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    set_run(p.add_run(f"{label}: "), bold=True)
    set_run(p.add_run(value))


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run(p.add_run(text))
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    set_run(p.add_run(f"{label}  "), bold=True, color=BLUE)
    set_run(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, RGBColor(31, 77, 120)),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("TEZCON EUROPE  |  DISCUSSION DRAFT"), size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Non-binding collaboration framework  |  29 July 2026"), size=8.5, color=MUTED)


def add_signature_block(doc):
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [4560, 4560], indent_dxa=120)
    labels = [
        ("For Ratata", "For TezCon Seattle"),
        ("Name: __________________________", "Name: __________________________"),
        ("Role: ___________________________", "Role: ___________________________"),
        ("Date: ___________________________", "Date: ___________________________"),
    ]
    for ridx, pair in enumerate(labels):
        for cidx, value in enumerate(pair):
            cell = table.cell(ridx, cidx)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            set_run(cell.paragraphs[0].add_run(value), bold=(ridx == 0))
            if ridx == 0:
                set_cell_shading(cell, LIGHT)


def build():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    even_odd = doc.settings._element.find(qn("w:evenAndOddHeaders"))
    if even_odd is not None:
        doc.settings._element.remove(even_odd)
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header_footer(section)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("MEMORANDUM OF UNDERSTANDING"), size=23, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_run(p.add_run("TezCon Europe · Frankfurt am Main"), size=14, color=MUTED)

    add_label_value(doc, "Between", "Ratata [insert full legal name and address]")
    add_label_value(doc, "And", "TezCon Seattle [insert organising entity or representative(s) and address]")
    add_label_value(doc, "Project", "TezCon Europe, proposed for Frankfurt am Main, Germany")
    add_label_value(doc, "Provisional timing", "October or November 2026; exact date to be confirmed")
    add_label_value(doc, "Document status", "Discussion draft — non-binding")

    add_callout(
        doc,
        "Purpose",
        "To record the shared intent, principles and working responsibilities for developing the first TezCon Europe in cooperation with the organisers of TezCon Seattle.",
    )

    doc.add_heading("1. Shared intention", level=1)
    add_body(
        doc,
        "Ratata and TezCon Seattle (together, the “Collaborators”) intend to explore and, subject to final agreement, jointly support a European edition of TezCon in Frankfurt am Main. The proposed event will carry forward TezCon’s grassroots character while responding to the needs and culture of the European Tezos community."
    )
    add_body(
        doc,
        "The working ambition is an accessible gathering for approximately 500 participants, bringing together artists, collectors, builders, musicians, organisers and people encountering Tezos for the first time."
    )

    doc.add_heading("2. Guiding principles", level=1)
    for item in (
        "From the community, for the community: program choices should prioritise participation, contribution and genuine connection.",
        "Continuity with TezCon: the European edition should respect the history, values and distinctive non-corporate spirit established in Seattle.",
        "Local authorship: the Frankfurt edition should be shaped with European communities rather than simply reproducing a US event.",
        "Artist and contributor care: fair treatment, clear expectations, accessibility and appropriate compensation should guide production decisions.",
        "Transparency: major decisions, funding relationships and conflicts of interest should be communicated openly between the Collaborators.",
        "Responsible growth: attendance and sponsor visibility should not be pursued at the expense of trust, safety or community character.",
    ):
        add_bullet(doc, item)

    doc.add_heading("3. Proposed event frame", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2450, 6910])
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(("Element", "Current working assumption")):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT)
        set_run(cell.paragraphs[0].add_run(text), bold=True)
    rows = (
        ("Location", "Frankfurt am Main, Germany"),
        ("Timing", "October or November 2026; 24–25 October currently appears in provisional Ratata project materials and remains subject to confirmation"),
        ("Scale", "Venue capacity around 500; attendance and registration targets to be agreed after venue and budget validation"),
        ("Format", "Opening gathering and exhibition, followed by a main day of talks, panels, workshops, art, music, minting and community exchange"),
        ("Audience", "European and international Tezos community, local creative and technology communities, students and newcomers"),
        ("Working language", "English, with selected German-language or bilingual elements where useful"),
    )
    for left, right in rows:
        cells = table.add_row().cells
        set_run(cells[0].paragraphs[0].add_run(left), bold=True)
        set_run(cells[1].paragraphs[0].add_run(right))
    set_table_geometry(table, [2450, 6910])

    doc.add_heading("4. Proposed responsibilities", level=1)
    doc.add_heading("4.1 Ratata", level=2)
    for item in (
        "Lead local production in Frankfurt, including venue sourcing, suppliers, permits, safety planning, insurance and on-site delivery.",
        "Develop the European program and local partnerships in consultation with TezCon Seattle.",
        "Prepare the working budget, fundraising materials, grant applications and sponsor proposals.",
        "Manage participant registration, local communications, contracting and payment administration through the agreed legal entity.",
        "Provide regular written production and financial updates to the joint steering group.",
    ):
        add_bullet(doc, item)

    doc.add_heading("4.2 TezCon Seattle", level=2)
    for item in (
        "Advise on TezCon’s history, values, community relationships and format.",
        "Confirm the conditions under which the TezCon name, identity and related materials may be used for the European edition.",
        "Support introductions to Tezos Commons, previous collaborators, contributors and potential ecosystem funders where appropriate.",
        "Participate in program review and help identify Seattle-based contributors or content that can create continuity between editions.",
        "Review public claims describing the relationship between TezCon Seattle and TezCon Europe before announcement.",
    ):
        add_bullet(doc, item)

    doc.add_heading("4.3 Shared responsibilities", level=2)
    for item in (
        "Agree the project narrative, values, event scope, program balance and public announcement.",
        "Establish a safe, respectful and inclusive participant environment, supported by a code of conduct.",
        "Protect the independence and trust of the community when evaluating sponsors or commercial activations.",
        "Document decisions, risks, commitments and lessons for a post-event review.",
    ):
        add_bullet(doc, item)

    doc.add_heading("5. Governance and decisions", level=1)
    add_body(
        doc,
        "Each Collaborator will nominate up to two representatives to a joint steering group. The steering group is expected to meet regularly during development and more frequently in the final production period."
    )
    add_body(doc, "The following matters require written agreement from both Collaborators before action:")
    for item in (
        "public confirmation of the event and use of the TezCon name;",
        "final event date, venue and public capacity;",
        "event identity and material changes to TezCon brand presentation;",
        "headline program, anchor sponsors and naming-rights arrangements;",
        "material budget changes or financial commitments involving both Collaborators;",
        "cancellation, postponement or material change of format.",
    ):
        add_bullet(doc, item)
    add_body(
        doc,
        "Day-to-day local production decisions may be made by Ratata within the scope and budget later approved by the steering group."
    )

    doc.add_heading("6. Name, brand and intellectual property", level=1)
    add_body(
        doc,
        "No transfer of ownership is intended by this memorandum. Each Collaborator retains ownership of its pre-existing names, logos, artwork, documents, systems and other materials."
    )
    add_body(
        doc,
        "Use of the TezCon name or identity for the European edition remains subject to written approval and an agreed brand arrangement. The Collaborators will decide how newly created event assets, recordings, photographs, program materials and domain or social-media accounts may be used after the event."
    )
    add_callout(
        doc,
        "Open decision",
        "Confirm whether “TezCon Europe” is an official regional edition, a licensed collaboration, or another defined relationship—and identify who can authorise the name on behalf of TezCon Seattle.",
    )

    doc.add_heading("7. Finance, fundraising and contracting", level=1)
    add_body(
        doc,
        "The Collaborators will develop a transparent project budget and funding plan. No Collaborator may make a financial commitment on behalf of the other without prior written authorisation."
    )
    for item in (
        "A single agreed legal entity should receive event income, sign production contracts, maintain accounting records and carry the primary financial risk.",
        "Grant applications and sponsor proposals must accurately describe the relationship between the Collaborators and identify the responsible contracting entity.",
        "Sponsor acceptance should be assessed against community values, reputational risk and program independence.",
        "The Collaborators should agree in writing how any deficit, surplus, cancelled-event costs, restricted grant funds and non-cash contributions will be treated.",
        "A post-event financial summary should be shared between the Collaborators and with funders where required.",
    ):
        add_bullet(doc, item)

    doc.add_heading("8. Communications, data and documentation", level=1)
    for item in (
        "Major public announcements and descriptions of the collaboration will be mutually reviewed before release.",
        "Press, sponsor and participant communications should use agreed terminology and clearly identify the responsible organiser.",
        "Registration and mailing-list data will be collected and processed only through agreed systems and in accordance with applicable data-protection law, including the GDPR.",
        "Photography, filming, livestreaming and reuse permissions will be addressed through appropriate participant notices and contributor agreements.",
        "Confidential information shared for planning or fundraising should not be disclosed outside the project team without permission, except where legally required.",
    ):
        add_bullet(doc, item)

    doc.add_heading("9. Safety, conduct and accessibility", level=1)
    add_body(
        doc,
        "Ratata will lead locally applicable safety, venue, insurance and emergency planning. The Collaborators intend to adopt a public code of conduct, a reporting process and clearly identified on-site contacts."
    )
    add_body(
        doc,
        "Venue and program decisions should consider physical accessibility, affordability, dietary requirements, sensory needs and routes for community members who require travel support."
    )

    doc.add_heading("10. Term, change and withdrawal", level=1)
    add_body(
        doc,
        "This memorandum takes effect when acknowledged by both Collaborators and remains a working framework until replaced by a definitive collaboration agreement, the project is concluded, or either Collaborator gives written notice that it no longer wishes to proceed."
    )
    add_body(
        doc,
        "If either Collaborator withdraws, both will work in good faith to minimise confusion, protect participants and funders, settle authorised commitments, and agree how the TezCon name and project materials may continue to be used."
    )

    doc.add_heading("11. Status of this memorandum", level=1)
    add_body(
        doc,
        "This document records present intentions and is not intended to create a partnership, joint venture, agency relationship, exclusivity obligation or legally binding duty to fund or deliver the event. Any binding commitments—including brand permission, financial responsibility, services, licensing, confidentiality, data processing or cancellation terms—should be set out in separately reviewed and signed agreements."
    )
    add_body(
        doc,
        "Each Collaborator remains responsible for obtaining its own legal, tax and insurance advice. This discussion draft is not legal advice."
    )

    doc.add_heading("12. Acknowledgement", level=1)
    add_body(
        doc,
        "By signing below, the representatives acknowledge that this memorandum accurately reflects the basis on which the Collaborators wish to continue discussions."
    )
    add_signature_block(doc)

    doc.add_page_break()
    doc.add_heading("Appendix A — Decisions required before public launch", level=1)
    decisions = (
        ("Parties", "Full legal name, address and authorised representative for Ratata and TezCon Seattle"),
        ("Relationship", "Official edition, licensed collaboration or another defined model"),
        ("Date", "Final October/November 2026 date; confirm or replace provisional 24–25 October"),
        ("Venue", "Venue, capacity, accessibility, production limitations and cancellation terms"),
        ("Legal organiser", "Entity signing contracts, selling tickets and receiving grants or sponsorship"),
        ("Budget authority", "Approved budget, approval thresholds, deficit exposure and treatment of surplus"),
        ("Insurance", "Public liability, event cancellation, equipment and other relevant cover"),
        ("Brand", "Name permission, visual identity, domains, social accounts and approval process"),
        ("Program", "Curatorial responsibility, contributor selection and compensation principles"),
        ("Sponsors", "Acceptance criteria, benefits, exclusions and approval rights"),
        ("Tickets", "Free, paid or deposit model; allocations for contributors, students and supported attendees"),
        ("Data", "Registration platform, controller/processor roles, privacy notice and mailing-list use"),
        ("Media", "Recording, livestream, photography, contributor releases and archive ownership"),
        ("Future editions", "Whether and how this collaboration affects later TezCon Europe editions"),
    )
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2450, 6910])
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(("Decision", "What must be agreed")):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT)
        set_run(cell.paragraphs[0].add_run(text), bold=True)
    for left, right in decisions:
        cells = table.add_row().cells
        set_run(cells[0].paragraphs[0].add_run(left), bold=True)
        set_run(cells[1].paragraphs[0].add_run(right))
    set_table_geometry(table, [2450, 6910])

    doc.add_heading("Appendix B — Recommended next steps", level=1)
    for item in (
        "Identify the authorised representatives and hold a focused alignment call using Appendix A as the agenda.",
        "Confirm the relationship model and written permission to develop and announce TezCon Europe.",
        "Agree a date window and request three comparable Frankfurt venue proposals.",
        "Name the legal and financial organiser before submitting grants or sponsor offers.",
        "Convert this memorandum into a short definitive collaboration agreement once the open decisions are resolved.",
    ):
        add_number(doc, item)

    doc.add_heading("Alignment call record", level=1)
    for label, value in (
        ("Date", "____________________________________________________________"),
        ("Participants", "______________________________________________________"),
        ("Decisions confirmed", "________________________________________________"),
        ("Actions and owners", "__________________________________________________"),
        ("Next review", "_______________________________________________________"),
    ):
        add_label_value(doc, label, value)

    doc.core_properties.title = "TezCon Europe — Memorandum of Understanding"
    doc.core_properties.subject = "Non-binding collaboration framework between Ratata and TezCon Seattle"
    doc.core_properties.author = "Ratata"
    doc.core_properties.keywords = "TezCon Europe, Ratata, TezCon Seattle, Frankfurt, memorandum of understanding"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
